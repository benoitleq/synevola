import os
import io
import time
import base64
import json
from pathlib import Path
from typing import Dict, List, Tuple, Optional
from datetime import datetime

import streamlit as st
import numpy as np
import requests
import soundfile as sf

# ==============================
# Optional heavy deps (tokenizers/docx)
# ==============================
HAVE_DOCX = True
try:
    from docx import Document
except Exception:
    HAVE_DOCX = False

# Torch pour afficher l'état CUDA
try:
    import torch
    CUDA_AVAILABLE = torch.cuda.is_available()
except Exception:
    CUDA_AVAILABLE = False

# Audio recorder component
HAVE_AUDIO_RECORDER = False
try:
    from audiorecorder import audiorecorder
    HAVE_AUDIO_RECORDER = True
except ImportError:
    try:
        from streamlit_mic_recorder import mic_recorder
        HAVE_AUDIO_RECORDER = "mic_recorder"
    except ImportError:
        HAVE_AUDIO_RECORDER = False

# Import du pipeline utilisateur
try:
    import audio_processing
    from audio_processing import process_audio, clean_temp_files
except Exception as e:
    st.error(f"Impossible d'importer 'audio_processing.py' : {e}")
    st.stop()


# ==============================
# Custom Audio Recorder Component with Mic Selection
# ==============================
def custom_audio_recorder_html(key: str = "custom_recorder"):
    """
    Composant d'enregistrement audio personnalisé avec:
    - Sélection du microphone
    - VU-mètre en temps réel
    - Contrôle de gain
    - Pause/Resume
    - Bouton de téléchargement direct
    """
    
    recorder_html = """
    <div id="recorder-container-{key}" style="font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; padding: 20px; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); border-radius: 15px; margin: 10px 0; box-shadow: 0 10px 40px rgba(0,0,0,0.2);">
        
        <div style="background: white; border-radius: 12px; padding: 20px;">
            
            <!-- Sélection du microphone -->
            <div style="margin-bottom: 20px;">
                <label style="display: block; margin-bottom: 8px; font-weight: 600; color: #333; font-size: 14px;">
                    🎤 Microphone
                </label>
                <select id="mic-select-{key}" style="width: 100%; padding: 12px; border-radius: 8px; border: 2px solid #e0e0e0; font-size: 14px; background: #fafafa; cursor: pointer; transition: border-color 0.3s;">
                    <option value="">Chargement des microphones...</option>
                </select>
            </div>
            
            <!-- Contrôle du gain -->
            <div style="margin-bottom: 20px;">
                <label style="display: block; margin-bottom: 8px; font-weight: 600; color: #333; font-size: 14px;">
                    🔊 Gain d'entrée : <span id="gain-value-{key}" style="color: #667eea; font-weight: bold;">100</span>%
                </label>
                <input type="range" id="gain-slider-{key}" min="0" max="200" value="100" 
                       style="width: 100%; cursor: pointer; height: 8px; -webkit-appearance: none; background: linear-gradient(90deg, #667eea, #764ba2); border-radius: 4px;">
            </div>
            
            <!-- VU-mètre -->
            <div style="margin-bottom: 20px;">
                <label style="display: block; margin-bottom: 8px; font-weight: 600; color: #333; font-size: 14px;">
                    📊 Niveau audio
                </label>
                <div style="background: #2d2d2d; border-radius: 8px; height: 30px; overflow: hidden; position: relative; box-shadow: inset 0 2px 10px rgba(0,0,0,0.3);">
                    <div id="vu-meter-{key}" style="height: 100%; width: 0%; background: linear-gradient(90deg, #00c853, #64dd17, #ffeb3b, #ff9800, #f44336); transition: width 0.05s ease-out; border-radius: 8px;"></div>
                    <div id="vu-peak-{key}" style="position: absolute; top: 0; left: 0; height: 100%; width: 3px; background: #ff1744; border-radius: 2px; transition: left 0.05s;"></div>
                </div>
                <div style="display: flex; justify-content: space-between; font-size: 11px; color: #888; margin-top: 4px; padding: 0 5px;">
                    <span>-60dB</span>
                    <span>-40dB</span>
                    <span>-20dB</span>
                    <span>-10dB</span>
                    <span>0dB</span>
                </div>
            </div>
            
            <!-- Boutons de contrôle -->
            <div style="display: flex; gap: 12px; margin-bottom: 20px;">
                <button id="start-btn-{key}" onclick="startRecording_{key}()" 
                        style="flex: 1; padding: 14px 20px; background: linear-gradient(135deg, #00c853, #00e676); color: white; border: none; border-radius: 10px; cursor: pointer; font-size: 15px; font-weight: 600; box-shadow: 0 4px 15px rgba(0,200,83,0.4); transition: transform 0.2s, box-shadow 0.2s;">
                    🔴 Démarrer
                </button>
                <button id="pause-btn-{key}" onclick="togglePause_{key}()" disabled
                        style="flex: 1; padding: 14px 20px; background: linear-gradient(135deg, #ff9800, #ffc107); color: white; border: none; border-radius: 10px; cursor: pointer; font-size: 15px; font-weight: 600; opacity: 0.5; transition: transform 0.2s, box-shadow 0.2s;">
                    ⏸️ Pause
                </button>
                <button id="stop-btn-{key}" onclick="stopRecording_{key}()" disabled
                        style="flex: 1; padding: 14px 20px; background: linear-gradient(135deg, #f44336, #e91e63); color: white; border: none; border-radius: 10px; cursor: pointer; font-size: 15px; font-weight: 600; opacity: 0.5; transition: transform 0.2s, box-shadow 0.2s;">
                    ⏹️ Arrêter
                </button>
            </div>
            
            <!-- Status -->
            <div id="status-{key}" style="text-align: center; padding: 12px 20px; background: linear-gradient(135deg, #e3f2fd, #bbdefb); border-radius: 10px; margin-bottom: 15px; font-weight: 500; color: #1565c0;">
                ⏳ Prêt à enregistrer
            </div>
            
            <!-- Timer -->
            <div id="timer-{key}" style="text-align: center; font-size: 36px; font-weight: 700; color: #333; margin-bottom: 20px; font-family: 'SF Mono', 'Monaco', 'Inconsolata', monospace; letter-spacing: 2px;">
                00:00:00
            </div>
            
            <!-- Audio preview et téléchargement -->
            <div id="preview-container-{key}" style="display: none; margin-top: 20px; padding: 20px; background: linear-gradient(135deg, #e8f5e9, #c8e6c9); border-radius: 12px;">
                <label style="display: block; margin-bottom: 10px; font-weight: 600; color: #2e7d32; font-size: 14px;">
                    ✅ Enregistrement terminé !
                </label>
                <audio id="audio-preview-{key}" controls style="width: 100%; margin-bottom: 15px; border-radius: 8px;"></audio>
                
                <!-- Bouton de téléchargement -->
                <a id="download-btn-{key}" download="enregistrement.wav" 
                   style="display: block; width: 100%; padding: 14px 20px; background: linear-gradient(135deg, #667eea, #764ba2); color: white; border: none; border-radius: 10px; cursor: pointer; font-size: 15px; font-weight: 600; text-align: center; text-decoration: none; box-shadow: 0 4px 15px rgba(102,126,234,0.4); transition: transform 0.2s;">
                    💾 Télécharger l'enregistrement (WAV)
                </a>
                
                <p style="margin-top: 15px; font-size: 13px; color: #558b2f; text-align: center;">
                    📁 Téléchargez le fichier puis importez-le dans l'onglet "Importer un fichier"
                </p>
            </div>
            
        </div>
    </div>
    
    <script>
    (function() {
        const key = '{key}';
        let mediaRecorder = null;
        let audioChunks = [];
        let audioContext = null;
        let analyser = null;
        let gainNode = null;
        let sourceNode = null;
        let mediaStream = null;
        let isRecording = false;
        let isPaused = false;
        let startTime = null;
        let pausedTime = 0;
        let pauseStartTime = null;
        let timerInterval = null;
        let peakLevel = 0;
        let animationId = null;
        
        // Éléments DOM
        const micSelect = document.getElementById('mic-select-' + key);
        const gainSlider = document.getElementById('gain-slider-' + key);
        const gainValue = document.getElementById('gain-value-' + key);
        const vuMeter = document.getElementById('vu-meter-' + key);
        const vuPeak = document.getElementById('vu-peak-' + key);
        const startBtn = document.getElementById('start-btn-' + key);
        const pauseBtn = document.getElementById('pause-btn-' + key);
        const stopBtn = document.getElementById('stop-btn-' + key);
        const status = document.getElementById('status-' + key);
        const timer = document.getElementById('timer-' + key);
        const previewContainer = document.getElementById('preview-container-' + key);
        const audioPreview = document.getElementById('audio-preview-' + key);
        const downloadBtn = document.getElementById('download-btn-' + key);
        
        // Charger la liste des microphones
        async function loadMicrophones() {
            try {
                await navigator.mediaDevices.getUserMedia({ audio: true });
                const devices = await navigator.mediaDevices.enumerateDevices();
                const audioInputs = devices.filter(d => d.kind === 'audioinput');
                
                micSelect.innerHTML = '';
                audioInputs.forEach((device, index) => {
                    const option = document.createElement('option');
                    option.value = device.deviceId;
                    option.text = device.label || 'Microphone ' + (index + 1);
                    micSelect.appendChild(option);
                });
                
                if (audioInputs.length === 0) {
                    micSelect.innerHTML = '<option value="">Aucun microphone détecté</option>';
                }
            } catch (err) {
                console.error('Erreur accès microphones:', err);
                micSelect.innerHTML = '<option value="">Erreur: ' + err.message + '</option>';
                status.textContent = '❌ Erreur: ' + err.message;
                status.style.background = 'linear-gradient(135deg, #ffebee, #ffcdd2)';
                status.style.color = '#c62828';
            }
        }
        
        // Mettre à jour le gain
        gainSlider.addEventListener('input', function() {
            gainValue.textContent = this.value;
            if (gainNode) {
                gainNode.gain.value = this.value / 100;
            }
        });
        
        // Mise à jour du VU-mètre
        function updateVuMeter() {
            if (!analyser || !isRecording) {
                vuMeter.style.width = '0%';
                return;
            }
            
            if (isPaused) {
                animationId = requestAnimationFrame(updateVuMeter);
                return;
            }
            
            const dataArray = new Uint8Array(analyser.frequencyBinCount);
            analyser.getByteFrequencyData(dataArray);
            
            let sum = 0;
            for (let i = 0; i < dataArray.length; i++) {
                sum += dataArray[i];
            }
            const average = sum / dataArray.length;
            const level = (average / 255) * 100;
            
            vuMeter.style.width = level + '%';
            
            if (level > peakLevel) {
                peakLevel = level;
            }
            vuPeak.style.left = peakLevel + '%';
            
            peakLevel = Math.max(0, peakLevel - 0.3);
            
            animationId = requestAnimationFrame(updateVuMeter);
        }
        
        // Mise à jour du timer
        function updateTimer() {
            if (!startTime || isPaused) return;
            
            const elapsed = (Date.now() - startTime - pausedTime) / 1000;
            const hours = Math.floor(elapsed / 3600);
            const minutes = Math.floor((elapsed % 3600) / 60);
            const seconds = Math.floor(elapsed % 60);
            
            timer.textContent = 
                String(hours).padStart(2, '0') + ':' +
                String(minutes).padStart(2, '0') + ':' +
                String(seconds).padStart(2, '0');
        }
        
        // Convertir AudioBuffer en WAV
        function audioBufferToWav(buffer) {
            const numChannels = 1;
            const sampleRate = buffer.sampleRate;
            const format = 1; // PCM
            const bitDepth = 16;
            
            const data = buffer.getChannelData(0);
            const dataLength = data.length * (bitDepth / 8);
            const headerLength = 44;
            const totalLength = headerLength + dataLength;
            
            const arrayBuffer = new ArrayBuffer(totalLength);
            const view = new DataView(arrayBuffer);
            
            // WAV Header
            const writeString = (offset, string) => {
                for (let i = 0; i < string.length; i++) {
                    view.setUint8(offset + i, string.charCodeAt(i));
                }
            };
            
            writeString(0, 'RIFF');
            view.setUint32(4, totalLength - 8, true);
            writeString(8, 'WAVE');
            writeString(12, 'fmt ');
            view.setUint32(16, 16, true);
            view.setUint16(20, format, true);
            view.setUint16(22, numChannels, true);
            view.setUint32(24, sampleRate, true);
            view.setUint32(28, sampleRate * numChannels * (bitDepth / 8), true);
            view.setUint16(32, numChannels * (bitDepth / 8), true);
            view.setUint16(34, bitDepth, true);
            writeString(36, 'data');
            view.setUint32(40, dataLength, true);
            
            // Audio data
            let offset = 44;
            for (let i = 0; i < data.length; i++) {
                const sample = Math.max(-1, Math.min(1, data[i]));
                view.setInt16(offset, sample < 0 ? sample * 0x8000 : sample * 0x7FFF, true);
                offset += 2;
            }
            
            return new Blob([arrayBuffer], { type: 'audio/wav' });
        }
        
        // Démarrer l'enregistrement
        window['startRecording_' + key] = async function() {
            try {
                const deviceId = micSelect.value;
                const constraints = {
                    audio: deviceId ? { 
                        deviceId: { exact: deviceId },
                        echoCancellation: false,
                        noiseSuppression: false,
                        autoGainControl: false
                    } : {
                        echoCancellation: false,
                        noiseSuppression: false,
                        autoGainControl: false
                    }
                };
                
                mediaStream = await navigator.mediaDevices.getUserMedia(constraints);
                
                audioContext = new (window.AudioContext || window.webkitAudioContext)();
                sourceNode = audioContext.createMediaStreamSource(mediaStream);
                
                gainNode = audioContext.createGain();
                gainNode.gain.value = gainSlider.value / 100;
                
                analyser = audioContext.createAnalyser();
                analyser.fftSize = 256;
                
                sourceNode.connect(gainNode);
                gainNode.connect(analyser);
                
                // Enregistrer avec le gain appliqué
                const dest = audioContext.createMediaStreamDestination();
                gainNode.connect(dest);
                
                // Utiliser un format compatible
                let mimeType = 'audio/webm;codecs=opus';
                if (!MediaRecorder.isTypeSupported(mimeType)) {
                    mimeType = 'audio/webm';
                }
                if (!MediaRecorder.isTypeSupported(mimeType)) {
                    mimeType = 'audio/ogg';
                }
                
                mediaRecorder = new MediaRecorder(dest.stream, { mimeType });
                audioChunks = [];
                
                mediaRecorder.ondataavailable = (event) => {
                    if (event.data.size > 0) {
                        audioChunks.push(event.data);
                    }
                };
                
                mediaRecorder.onstop = async () => {
                    // Créer le blob audio
                    const audioBlob = new Blob(audioChunks, { type: mimeType });
                    
                    // Convertir en WAV pour meilleure compatibilité
                    try {
                        const arrayBuffer = await audioBlob.arrayBuffer();
                        const audioCtx = new (window.AudioContext || window.webkitAudioContext)();
                        const audioBuffer = await audioCtx.decodeAudioData(arrayBuffer);
                        const wavBlob = audioBufferToWav(audioBuffer);
                        
                        const wavUrl = URL.createObjectURL(wavBlob);
                        audioPreview.src = wavUrl;
                        downloadBtn.href = wavUrl;
                        
                        // Nom du fichier avec timestamp
                        const now = new Date();
                        const filename = 'enregistrement_' + 
                            now.getFullYear() + 
                            String(now.getMonth() + 1).padStart(2, '0') +
                            String(now.getDate()).padStart(2, '0') + '_' +
                            String(now.getHours()).padStart(2, '0') +
                            String(now.getMinutes()).padStart(2, '0') +
                            String(now.getSeconds()).padStart(2, '0') + '.wav';
                        downloadBtn.download = filename;
                        
                    } catch (e) {
                        // Fallback: utiliser le format original
                        console.warn('Conversion WAV échouée, utilisation du format original:', e);
                        const audioUrl = URL.createObjectURL(audioBlob);
                        audioPreview.src = audioUrl;
                        downloadBtn.href = audioUrl;
                        downloadBtn.download = 'enregistrement.webm';
                    }
                    
                    previewContainer.style.display = 'block';
                };
                
                mediaRecorder.start(100);
                isRecording = true;
                isPaused = false;
                startTime = Date.now();
                pausedTime = 0;
                
                // UI updates
                startBtn.disabled = true;
                startBtn.style.opacity = '0.5';
                pauseBtn.disabled = false;
                pauseBtn.style.opacity = '1';
                pauseBtn.style.boxShadow = '0 4px 15px rgba(255,152,0,0.4)';
                stopBtn.disabled = false;
                stopBtn.style.opacity = '1';
                stopBtn.style.boxShadow = '0 4px 15px rgba(244,67,54,0.4)';
                micSelect.disabled = true;
                status.textContent = '🔴 Enregistrement en cours...';
                status.style.background = 'linear-gradient(135deg, #ffebee, #ffcdd2)';
                status.style.color = '#c62828';
                previewContainer.style.display = 'none';
                
                updateVuMeter();
                timerInterval = setInterval(updateTimer, 100);
                
            } catch (err) {
                console.error('Erreur démarrage:', err);
                status.textContent = '❌ Erreur: ' + err.message;
                status.style.background = 'linear-gradient(135deg, #ffebee, #ffcdd2)';
                status.style.color = '#c62828';
            }
        };
        
        // Pause/Resume
        window['togglePause_' + key] = function() {
            if (!mediaRecorder) return;
            
            if (isPaused) {
                mediaRecorder.resume();
                isPaused = false;
                pausedTime += Date.now() - pauseStartTime;
                pauseBtn.innerHTML = '⏸️ Pause';
                status.textContent = '🔴 Enregistrement en cours...';
                status.style.background = 'linear-gradient(135deg, #ffebee, #ffcdd2)';
                status.style.color = '#c62828';
                timerInterval = setInterval(updateTimer, 100);
            } else {
                mediaRecorder.pause();
                isPaused = true;
                pauseStartTime = Date.now();
                clearInterval(timerInterval);
                pauseBtn.innerHTML = '▶️ Reprendre';
                status.textContent = '⏸️ En pause';
                status.style.background = 'linear-gradient(135deg, #fff3e0, #ffe0b2)';
                status.style.color = '#e65100';
            }
        };
        
        // Arrêter l'enregistrement
        window['stopRecording_' + key] = function() {
            if (!mediaRecorder) return;
            
            isRecording = false;
            mediaRecorder.stop();
            
            if (mediaStream) {
                mediaStream.getTracks().forEach(track => track.stop());
            }
            
            if (animationId) {
                cancelAnimationFrame(animationId);
            }
            
            if (timerInterval) {
                clearInterval(timerInterval);
            }
            
            // UI updates
            startBtn.disabled = false;
            startBtn.style.opacity = '1';
            pauseBtn.disabled = true;
            pauseBtn.style.opacity = '0.5';
            pauseBtn.style.boxShadow = 'none';
            pauseBtn.innerHTML = '⏸️ Pause';
            stopBtn.disabled = true;
            stopBtn.style.opacity = '0.5';
            stopBtn.style.boxShadow = 'none';
            micSelect.disabled = false;
            status.textContent = '✅ Enregistrement terminé !';
            status.style.background = 'linear-gradient(135deg, #e8f5e9, #c8e6c9)';
            status.style.color = '#2e7d32';
            vuMeter.style.width = '0%';
            vuPeak.style.left = '0%';
        };
        
        // Initialisation
        loadMicrophones();
        navigator.mediaDevices.addEventListener('devicechange', loadMicrophones);
    })();
    </script>
    """.replace('{key}', key)
    
    return recorder_html


# ==============================
# LM Studio Status Check
# ==============================
def check_lmstudio_status(base_url: str) -> Tuple[bool, str, List[str]]:
    try:
        url = f"{base_url.rstrip('/')}/v1/models"
        resp = requests.get(url, timeout=5)
        if resp.status_code == 200:
            data = resp.json()
            models = [m.get("id") for m in data.get("data", []) if m.get("id")]
            if models:
                return True, f"Connecté ({len(models)} modèle(s))", models
            else:
                return True, "Connecté (aucun modèle chargé)", []
        else:
            return False, f"Erreur HTTP {resp.status_code}", []
    except requests.exceptions.ConnectionError:
        return False, "Non connecté", []
    except requests.exceptions.Timeout:
        return False, "Timeout", []
    except Exception as e:
        return False, f"Erreur: {str(e)[:50]}", []


# ==============================
# Tokenization utilities
# ==============================
@st.cache_resource(show_spinner=False)
def get_hf_tokenizer(model_id: str):
    try:
        from transformers import AutoTokenizer
        tok = AutoTokenizer.from_pretrained(model_id, use_fast=True)
        return tok
    except Exception as e:
        st.warning(f"Impossible de charger le tokenizer HF '{model_id}': {e}")
        return None

def encode_tokens(text: str, backend: str, hf_model_id: str = "Qwen/Qwen2-7B-Instruct") -> List[int]:
    if not text:
        return []
    if backend == "HuggingFace":
        tok = get_hf_tokenizer(hf_model_id)
        if tok is None:
            backend = "tiktoken"
        else:
            try:
                return tok.encode(text)
            except Exception as e:
                st.warning(f"Echec encodage HF, fallback tiktoken: {e}")
                backend = "tiktoken"
    try:
        import tiktoken
        enc = tiktoken.get_encoding("cl100k_base")
        return enc.encode(text)
    except Exception:
        return text.split()

def count_tokens(text: str, backend: str, hf_model_id: str) -> int:
    return len(encode_tokens(text, backend, hf_model_id))

def chunk_text_by_tokens(text: str, max_tokens: int, backend: str, hf_model_id: str, overlap: int = 200) -> List[str]:
    ids = encode_tokens(text, backend, hf_model_id)
    if not ids:
        return []
    chunks: List[List[int]] = []
    start = 0
    n = len(ids)
    while start < n:
        end = min(start + max_tokens, n)
        chunks.append(ids[start:end])
        if end >= n:
            break
        start = end - overlap if end - overlap > 0 else end
    decoded: List[str] = []
    if backend == "HuggingFace" and get_hf_tokenizer(hf_model_id) is not None:
        tok = get_hf_tokenizer(hf_model_id)
        for c in chunks:
            decoded.append(tok.decode(c))
    else:
        words = text.split()
        if len(ids) == 0:
            return [text]
        ratio = max(1, int(round(len(words) / len(ids))))
        idx = 0
        for c in chunks:
            take = max(1, len(c) * ratio)
            decoded.append(" ".join(words[idx: idx + take]))
            idx += take - min(overlap, 50)
    return decoded


# ==============================
# LM Studio helpers
# ==============================
def _lmstudio_base(base_url: str) -> str:
    return base_url.rstrip("/")

def _lmstudio_headers():
    key = os.environ.get("LMSTUDIO_API_KEY", "").strip()
    if key:
        return {"Authorization": f"Bearer {key}"}
    return {"Authorization": "Bearer lm-studio"}

def list_lmstudio_models(base_url: str) -> List[str]:
    try:
        url = f"{_lmstudio_base(base_url)}/v1/models"
        resp = requests.get(url, timeout=8, headers=_lmstudio_headers())
        resp.raise_for_status()
        data = resp.json()
        return [m.get("id") for m in data.get("data", []) if m.get("id")]
    except Exception:
        return []

def call_chat_completions(base_url: str, model: str, system_prompt: str, user_prompt: str,
                          temperature: float, max_tokens: int):
    url = f"{_lmstudio_base(base_url)}/v1/chat/completions"
    
    system_prompt_clean = system_prompt.strip() if system_prompt else ""
    user_prompt_clean = user_prompt.strip() if user_prompt else ""
    
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system_prompt_clean},
            {"role": "user", "content": user_prompt_clean},
        ],
        "temperature": float(temperature),
        "max_tokens": int(max_tokens),
        "stream": False
    }
    
    try:
        resp = requests.post(url, json=payload, timeout=300, headers=_lmstudio_headers())
        resp.raise_for_status()
        data = resp.json()
        if "choices" in data and data["choices"]:
            return data["choices"][0]["message"]["content"]
        return ""
    except requests.exceptions.HTTPError as e:
        error_detail = ""
        try:
            error_detail = resp.text
        except:
            pass
        raise requests.exceptions.HTTPError(f"{e} - Détail: {error_detail[:300]}")

def call_completions(base_url: str, model: str, prompt: str,
                     temperature: float, max_tokens: int):
    url = f"{_lmstudio_base(base_url)}/v1/completions"
    payload = {
        "model": model,
        "prompt": prompt,
        "temperature": float(temperature),
        "max_tokens": int(max_tokens),
        "stream": False
    }
    resp = requests.post(url, json=payload, timeout=300, headers=_lmstudio_headers())
    resp.raise_for_status()
    data = resp.json()
    if "choices" in data and data["choices"]:
        return data["choices"][0].get("text", "")
    return ""

def chat_or_complete(base_url: str, model: str, system_prompt: str, user_prompt: str,
                     temperature: float = 0.2, max_tokens: int = 512, mode: str = "auto") -> str:
    last_error = None
    
    if mode in ("auto", "chat"):
        try:
            result = call_chat_completions(base_url, model, system_prompt, user_prompt, temperature, max_tokens)
            if result:
                return result
        except requests.exceptions.HTTPError as e:
            last_error = e
            if mode == "chat":
                raise ValueError(f"Erreur API chat/completions: {e}.")
        except Exception as e:
            last_error = e
            if mode == "chat":
                raise
    
    if mode in ("auto", "completions"):
        try:
            merged_prompt = f"### System:\n{system_prompt}\n\n### User:\n{user_prompt}\n\n### Assistant:\n"
            result = call_completions(base_url, model, merged_prompt, temperature, max_tokens)
            if result:
                return result
        except requests.exceptions.HTTPError as e:
            raise ValueError(
                f"Échec des deux endpoints API.\n"
                f"- /v1/chat/completions: {last_error}\n"
                f"- /v1/completions: {e}\n\n"
                f"Vérifiez LM Studio et le Context Length."
            )
        except Exception as e:
            if last_error:
                raise ValueError(f"Échec chat ({last_error}) et completions ({e})")
            raise
    
    if last_error:
        raise last_error
    return ""


# ==============================
# Utils transcription
# ==============================
def format_duration(seconds: float) -> str:
    minutes = int(seconds // 60)
    remaining_seconds = int(seconds % 60)
    return f"{minutes}m{remaining_seconds}s"

def format_duration_full(seconds: float) -> str:
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    if hours > 0:
        return f"{hours}h{minutes:02d}m{secs:02d}s"
    return f"{minutes}m{secs:02d}s"

def normalize_audio(in_path: str, out_path: str) -> str:
    data, sr = sf.read(in_path)
    if len(data.shape) > 1:
        data = np.mean(data, axis=1)
    max_amp = np.max(np.abs(data))
    if max_amp > 0:
        data = data / max_amp
    sf.write(out_path, data, sr)
    return out_path

def join_transcription_text(maybe_list_or_str) -> str:
    if maybe_list_or_str is None:
        return ""
    if isinstance(maybe_list_or_str, (list, tuple)):
        return " ".join([str(x).strip() for x in maybe_list_or_str if str(x).strip()])
    return str(maybe_list_or_str).strip()

def ensure_dir(path: str) -> None:
    Path(path).mkdir(parents=True, exist_ok=True)

def overwrite_pyannote_config_if_provided(custom_config_path: Optional[str]):
    if not custom_config_path:
        return
    if not hasattr(audio_processing, "load_pipeline_from_pretrained"):
        return
    _orig_loader = audio_processing.load_pipeline_from_pretrained
    def _wrapped_loader(_ignored_input_path):
        return _orig_loader(custom_config_path)
    audio_processing.load_pipeline_from_pretrained = _wrapped_loader

def build_text_from_transcription(transcriptions, diarization_enabled: bool, include_speakers: bool = True) -> str:
    if not transcriptions:
        return ""
    if not diarization_enabled:
        return "\n".join([str(t) for t in transcriptions])
    lines = []
    for start, end, spk, txt in transcriptions:
        s = f"{format_duration(start)} - {format_duration(end)}"
        if include_speakers:
            lines.append(f"{s} — {spk}: {txt}")
        else:
            lines.append(f"{s} — {txt}")
    return "\n".join(lines)


# ==============================
# Fonction de résumé
# ==============================
def generate_summary(text_for_llm: str, lm_base_url: str, lm_model: str, 
                     system_prompt: str, user_prompt: str,
                     temperature: float, max_tokens: int, 
                     api_mode_key: str, token_backend: str, hf_tokenizer_model: str,
                     chunk_size: int, overlap: int,
                     mode_resume: str, status_container=None) -> Tuple[str, List[str]]:
    total_tokens = count_tokens(text_for_llm, token_backend, hf_tokenizer_model)
    chunk_summaries = []
    summary_text = ""
    
    if mode_resume.startswith("Résumé direct") or total_tokens <= chunk_size:
        prompt = f"{user_prompt}\n\nTexte à résumer:\n```text\n{text_for_llm}\n```"
        if status_container:
            status_container.write("📝 Génération du résumé direct...")
        summary_text = chat_or_complete(lm_base_url, lm_model, system_prompt, prompt,
                                       temperature=temperature, max_tokens=max_tokens, mode=api_mode_key)
    else:
        if status_container:
            status_container.write("✂️ Découpage du texte en blocs...")
        
        chunks = chunk_text_by_tokens(text_for_llm, int(chunk_size), token_backend, hf_tokenizer_model, overlap=int(overlap))
        
        if status_container:
            status_container.write(f"📚 {len(chunks)} blocs à résumer.")
        
        partial_summaries = []
        for i, ch in enumerate(chunks, 1):
            if status_container:
                status_container.write(f"🔄 Résumé bloc {i}/{len(chunks)}...")
            p = f"{user_prompt}\n\nTu résumes le **bloc {i}/{len(chunks)}** ci-dessous.\n```text\n{ch}\n```"
            smry = chat_or_complete(lm_base_url, lm_model, system_prompt, p,
                                    temperature=temperature, max_tokens=max_tokens, mode=api_mode_key)
            partial_summaries.append(f"### Bloc {i}\n{smry}")
        
        chunk_summaries = partial_summaries
        
        if status_container:
            status_container.write("🔗 Synthèse finale en cours...")
        
        joined = "\n\n".join(chunk_summaries)
        synth_prompt = (
            f"{user_prompt}\n\nVoici les résumés des blocs :\n{joined}\n\n"
            "Produis **un seul résumé final** structuré. Ne répète pas bloc par bloc."
        )
        summary_text = chat_or_complete(lm_base_url, lm_model, system_prompt, synth_prompt,
                                        temperature=temperature, max_tokens=max_tokens, mode=api_mode_key)
    
    return summary_text, chunk_summaries


# ==============================
# App
# ==============================
st.set_page_config(page_title="Synevola", layout="wide", page_icon="🩺")
st.title("🩺 Synevola")
st.caption("Transcription & Résumé aidé par IA en local")

# Workspace
workspace = Path("./streamlit_workspace")
ensure_dir(workspace)
recordings_dir = workspace / "recordings"
ensure_dir(recordings_dir)

# ==============================
# Sidebar
# ==============================
with st.sidebar:
    st.header("⚙️ Configuration")
    
    # Status indicators
    st.subheader("📊 État du système")
    
    default_base = os.environ.get("LMSTUDIO_BASE_URL", "http://127.0.0.1:1234")
    lm_base_url = st.text_input("URL LM Studio", value=default_base)
    
    col1, col2 = st.columns(2)
    with col1:
        if CUDA_AVAILABLE:
            st.success("CUDA ✅")
        else:
            st.error("CUDA ❌")
    
    with col2:
        lm_available, lm_status, lm_models_list = check_lmstudio_status(lm_base_url)
        if lm_available and lm_models_list:
            st.success("LM Studio ✅")
        elif lm_available:
            st.warning("LM Studio ⚠️")
        else:
            st.error("LM Studio ❌")
    
    st.caption(f"LM Studio: {lm_status}")
    
    if st.button("🔄 Rafraîchir statut"):
        st.rerun()
    
    st.divider()
    
    # STT Settings
    st.header("🎤 Speech-to-Text")
    model_choice = st.selectbox(
        "Modèle STT",
        ["tiny", "base", "small", "medium", "large", "faster-whisper"],
        index=2
    )
    diarization_enabled = st.toggle("Activer la diarisation", value=True)
    postproc_enabled = st.toggle("Normalisation audio (mono)", value=False)
    custom_pyannote_config = st.text_input("Config pyannote (optionnel)", value="")
    
    st.divider()
    
    # LLM Settings
    st.header("🧠 LLM (Résumé)")
    
    if "lm_models" not in st.session_state:
        st.session_state.lm_models = lm_models_list
    
    if st.button("🔄 Rafraîchir modèles"):
        st.session_state.lm_models = list_lmstudio_models(lm_base_url)
    
    if st.session_state.lm_models:
        lm_model = st.selectbox("Modèle LLM", st.session_state.lm_models, index=0)
    else:
        st.warning("Aucun modèle détecté")
        lm_model = st.text_input("Nom du modèle", value="mistral-7b-instruct-v0.3")
    
    api_mode = st.selectbox("Mode API", ["Auto (chat→completions)", "Chat seulement", "Completions seulement"], index=0)
    api_mode_key = {"Auto (chat→completions)": "auto", "Chat seulement": "chat", "Completions seulement": "completions"}[api_mode]
    
    default_ctx = 16384
    if lm_model:
        if "qwen" in lm_model.lower():
            default_ctx = 32768
        elif "mistral" in lm_model.lower():
            default_ctx = 8192
    context_limit = st.number_input("Contexte max (tokens)", min_value=1024, max_value=256000, value=default_ctx, step=1024)
    
    st.divider()
    
    # Tokenization
    st.header("🔢 Tokenization")
    token_backend = st.selectbox("Backend tokens", ["HuggingFace", "tiktoken (approx)"], index=0)
    hf_tokenizer_model = st.text_input("Tokenizer HF", value="Qwen/Qwen2-7B-Instruct")

# Session state
st.session_state.setdefault("transcriptions", [])
st.session_state.setdefault("speaker_files", {})
st.session_state.setdefault("speaker_mapping", {})
st.session_state.setdefault("audio_path", None)
st.session_state.setdefault("summary_text", "")
st.session_state.setdefault("chunk_summaries", [])

# ==============================
# Main content - Source audio selection
# ==============================
st.subheader("📁 Source audio")

source_tab1, source_tab2, source_tab3 = st.tabs(["📤 Importer un fichier", "🎙️ Enregistrement avancé", "🎤 Enregistrement simple"])

audio_path = None

with source_tab1:
    uploaded_file = st.file_uploader(
        "Charger un fichier audio", 
        type=["mp3", "wav", "ogg", "flac", "m4a", "webm"],
        help="Formats supportés: MP3, WAV, OGG, FLAC, M4A, WebM"
    )
    
    if uploaded_file is not None:
        raw_path = workspace / uploaded_file.name
        with open(raw_path, "wb") as f:
            f.write(uploaded_file.read())
        st.success(f"✅ Fichier chargé : {raw_path.name}")
        audio_path = str(raw_path)
        st.session_state.audio_path = audio_path
        
        with open(raw_path, "rb") as f:
            st.audio(f.read())

with source_tab2:
    st.markdown("### 🎙️ Enregistrement avancé")
    st.info("**Fonctionnalités** : Sélection du microphone • Réglage du gain • VU-mètre • Pause/Reprise")
    
    # Afficher le composant HTML personnalisé
    st.components.v1.html(custom_audio_recorder_html(key="adv_rec"), height=700, scrolling=True)
    
    st.warning("""
    **📥 Après l'enregistrement :**
    1. Cliquez sur **"💾 Télécharger l'enregistrement"** dans le composant ci-dessus
    2. Le fichier WAV sera téléchargé sur votre ordinateur
    3. Importez-le via l'onglet **"📤 Importer un fichier"**
    """)

with source_tab3:
    st.markdown("### 🎤 Enregistrement simple")
    
    if HAVE_AUDIO_RECORDER == True:
        st.info("Cliquez pour démarrer/arrêter. Utilisez ⏸️ pour mettre en pause.")
        
        audio = audiorecorder(
            start_prompt="🔴 Démarrer",
            stop_prompt="⏹️ Arrêter", 
            pause_prompt="⏸️ Pause",
            show_visualizer=True,
            key="simple_audio_recorder"
        )
        
        if len(audio) > 0:
            duration_sec = len(audio) / 1000.0
            st.success(f"✅ Enregistrement : {format_duration_full(duration_sec)}")
            st.audio(audio.export().read())
            
            col1, col2 = st.columns(2)
            with col1:
                recording_name = st.text_input(
                    "Nom",
                    value=f"rec_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
                    key="rec_name_simple"
                )
            
            with col2:
                if st.button("💾 Sauvegarder", type="primary", use_container_width=True, key="save_simple"):
                    save_path = recordings_dir / f"{recording_name}.wav"
                    audio.export(str(save_path), format="wav")
                    st.session_state.audio_path = str(save_path)
                    st.success(f"✅ Sauvegardé!")
                    st.rerun()
    
    elif HAVE_AUDIO_RECORDER == "mic_recorder":
        st.info("Cliquez pour démarrer, cliquez à nouveau pour arrêter.")
        
        audio = mic_recorder(
            start_prompt="🔴 Démarrer",
            stop_prompt="⏹️ Arrêter",
            just_once=False,
            use_container_width=True,
            format="wav",
            key="mic_rec_simple"
        )
        
        if audio:
            st.success("✅ Enregistrement terminé")
            st.audio(audio['bytes'])
            
            col1, col2 = st.columns(2)
            with col1:
                recording_name = st.text_input(
                    "Nom",
                    value=f"rec_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
                    key="rec_name_mic"
                )
            
            with col2:
                if st.button("💾 Sauvegarder", type="primary", use_container_width=True, key="save_mic"):
                    save_path = recordings_dir / f"{recording_name}.wav"
                    with open(save_path, "wb") as f:
                        f.write(audio['bytes'])
                    st.session_state.audio_path = str(save_path)
                    st.success(f"✅ Sauvegardé!")
                    st.rerun()
    
    else:
        st.warning("""
        ⚠️ **Composant non installé**
        
        ```bash
        pip install streamlit-audiorecorder
        # ou
        pip install streamlit-mic-recorder
        ```
        """)
        
        try:
            audio_value = st.audio_input("🎤 Enregistrer", key="native_rec")
            
            if audio_value:
                st.audio(audio_value)
                
                col1, col2 = st.columns(2)
                with col1:
                    recording_name = st.text_input("Nom", value=f"rec_{datetime.now().strftime('%Y%m%d_%H%M%S')}", key="rec_name_native")
                
                with col2:
                    if st.button("💾 Sauvegarder", type="primary", use_container_width=True, key="save_native"):
                        save_path = recordings_dir / f"{recording_name}.wav"
                        with open(save_path, "wb") as f:
                            f.write(audio_value.getvalue())
                        st.session_state.audio_path = str(save_path)
                        st.rerun()
        except:
            st.error("Mettez à jour Streamlit: `pip install --upgrade streamlit`")

# Afficher l'audio sélectionné
if st.session_state.audio_path and os.path.exists(st.session_state.audio_path):
    st.info(f"🎵 Audio sélectionné : `{Path(st.session_state.audio_path).name}`")
    audio_path = st.session_state.audio_path

# Liste des enregistrements précédents
existing_recordings = list(recordings_dir.glob("*.wav")) + list(recordings_dir.glob("*.mp3"))
if existing_recordings:
    with st.expander(f"📂 Enregistrements précédents ({len(existing_recordings)})"):
        for rec in sorted(existing_recordings, reverse=True)[:10]:
            col1, col2, col3 = st.columns([3, 1, 1])
            with col1:
                st.text(rec.name)
            with col2:
                if st.button("▶️", key=f"use_{rec.name}"):
                    st.session_state.audio_path = str(rec)
                    st.rerun()
            with col3:
                if st.button("🗑️", key=f"del_{rec.name}"):
                    rec.unlink()
                    st.rerun()

# Update audio_path
if st.session_state.audio_path and os.path.exists(st.session_state.audio_path):
    audio_path = st.session_state.audio_path

# Pre-processing
current_audio_path = audio_path
if postproc_enabled and current_audio_path:
    try:
        normalized_path = str(workspace / "normalized_audio.wav")
        normalize_audio(current_audio_path, normalized_path)
        current_audio_path = normalized_path
    except Exception as e:
        st.warning(f"Échec normalisation : {e}")

# ==============================
# Paramètres de résumé
# ==============================
st.write("---")
st.subheader("📝 Paramètres de résumé")

with st.expander("Configurer les prompts et paramètres", expanded=False):
    default_system = (
        "You are a helpful medical summarization assistant. "
        "Write in clear French aimed at cardiologists. Use concise bullet points. "
        "Attribute speaker insights when relevant. Never invent facts."
    )
    system_prompt = st.text_area("System prompt", value=default_system, height=100)
    user_prompt = st.text_area(
        "Instructions utilisateur",
        value="Résume le texte ci-dessous pour un cardiologue. Structure en: 1) Contexte, 2) Points clés, 3) Décisions/Actions, 4) Questions ouvertes. Garde les chiffres importants, style concis.",
        height=100
    )
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        temperature = st.slider("Température", 0.0, 1.5, 0.2, 0.05)
    with col2:
        max_tokens_output = st.slider("Max tokens sortie", 128, 4096, 1024, 32)
    with col3:
        chunk_size = st.number_input("Taille bloc (tokens)", min_value=1000, max_value=128000, value=6000, step=500)
    with col4:
        overlap = st.number_input("Chevauchement", min_value=0, max_value=1000, value=200, step=50)
    
    mode_resume = st.selectbox("Mode de résumé", ["Résumé direct (1 bloc)", "Résumé par blocs + synthèse"], index=0)
    include_speakers = st.toggle("Inclure locuteurs/timestamps", value=True)

# ==============================
# Action buttons
# ==============================
st.write("---")

col_btn1, col_btn2, col_btn3 = st.columns([1, 1, 1])

with col_btn1:
    process_btn = st.button(
        "🎤 Transcrire uniquement", 
        type="secondary", 
        disabled=(current_audio_path is None),
        use_container_width=True
    )

with col_btn2:
    can_run_full = (current_audio_path is not None) and lm_available and lm_model
    full_process_btn = st.button(
        "🚀 Transcrire + Résumer", 
        type="primary", 
        disabled=(not can_run_full),
        use_container_width=True
    )

with col_btn3:
    clear_btn = st.button("🧹 Nettoyer", use_container_width=True)

if clear_btn:
    try:
        clean_temp_files()
        st.session_state.transcriptions = []
        st.session_state.summary_text = ""
        st.session_state.chunk_summaries = []
        st.success("Nettoyé!")
    except Exception as e:
        st.error(f"Erreur: {e}")

# ==============================
# Processing logic
# ==============================
run_transcription = process_btn or full_process_btn
run_summary_after = full_process_btn

if run_transcription and current_audio_path:
    with st.status("🎤 Transcription en cours...", expanded=True) as status:
        try:
            if custom_pyannote_config.strip():
                overwrite_pyannote_config_if_provided(custom_pyannote_config.strip())
            
            st.write(f"→ Modèle: **{model_choice}** | Diarisation: **{'Oui' if diarization_enabled else 'Non'}**")
            st.write(f"→ Fichier: `{Path(current_audio_path).name}`")
            
            transcriptions, speaker_files = process_audio(
                current_audio_path,
                diarization_enabled=diarization_enabled,
                model_name=model_choice
            )
            
            if not diarization_enabled:
                st.session_state.transcriptions = [str(t).strip() for t in (transcriptions or []) if str(t).strip()]
                st.session_state.speaker_files = {}
                st.session_state.speaker_mapping = {}
            else:
                cleaned = []
                unique_speakers = set()
                for item in transcriptions or []:
                    try:
                        start, end, spk, txt = item
                        unique_speakers.add(str(spk))
                        cleaned.append((float(start), float(end), str(spk), join_transcription_text(txt)))
                    except Exception:
                        cleaned.append((0.0, 0.0, "SPEAKER", join_transcription_text(item)))
                st.session_state.transcriptions = cleaned
                st.session_state.speaker_files = speaker_files or {}
                for s in unique_speakers:
                    st.session_state.speaker_mapping.setdefault(s, s)
            
            status.update(label="✅ Transcription terminée", state="complete")
            
        except Exception as e:
            status.update(label="❌ Erreur", state="error")
            st.error(f"Erreur: {e}")
            run_summary_after = False

if run_summary_after and st.session_state.transcriptions and lm_model:
    with st.status("🧠 Résumé...", expanded=True) as status:
        try:
            text_for_llm = build_text_from_transcription(
                st.session_state.transcriptions, 
                diarization_enabled, 
                include_speakers
            )
            
            total_tokens = count_tokens(text_for_llm, "HuggingFace" if token_backend.startswith("HuggingFace") else "tiktoken", hf_tokenizer_model)
            st.write(f"📊 {total_tokens} tokens")
            
            summary, chunks = generate_summary(
                text_for_llm=text_for_llm,
                lm_base_url=lm_base_url,
                lm_model=lm_model,
                system_prompt=system_prompt,
                user_prompt=user_prompt,
                temperature=temperature,
                max_tokens=max_tokens_output,
                api_mode_key=api_mode_key,
                token_backend="HuggingFace" if token_backend.startswith("HuggingFace") else "tiktoken",
                hf_tokenizer_model=hf_tokenizer_model,
                chunk_size=chunk_size,
                overlap=overlap,
                mode_resume=mode_resume,
                status_container=status
            )
            
            st.session_state.summary_text = summary
            st.session_state.chunk_summaries = chunks
            status.update(label="✅ Résumé généré", state="complete")
            
        except Exception as e:
            status.update(label="❌ Erreur", state="error")
            st.error(f"Erreur: {e}")

# ==============================
# Display results
# ==============================
st.write("---")

if st.session_state.transcriptions or st.session_state.summary_text:
    tab1, tab2, tab3 = st.tabs(["📝 Résumé", "🧾 Transcription", "📤 Export"])
    
    with tab1:
        if st.session_state.summary_text:
            st.markdown("### Résumé")
            st.markdown(st.session_state.summary_text)
            
            if st.button("🔄 Régénérer", key="regen"):
                with st.status("🧠 Régénération...", expanded=True) as status:
                    try:
                        text_for_llm = build_text_from_transcription(st.session_state.transcriptions, diarization_enabled, include_speakers)
                        summary, chunks = generate_summary(
                            text_for_llm, lm_base_url, lm_model, system_prompt, user_prompt,
                            temperature, max_tokens_output, api_mode_key,
                            "HuggingFace" if token_backend.startswith("HuggingFace") else "tiktoken",
                            hf_tokenizer_model, chunk_size, overlap, mode_resume, status
                        )
                        st.session_state.summary_text = summary
                        st.session_state.chunk_summaries = chunks
                        status.update(label="✅ OK", state="complete")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Erreur: {e}")
        else:
            st.info("Aucun résumé. Lancez le traitement complet.")
            if st.session_state.transcriptions and lm_available:
                if st.button("🧠 Générer le résumé", type="primary"):
                    # Même logique...
                    pass
    
    with tab2:
        if st.session_state.transcriptions:
            text_for_llm = build_text_from_transcription(st.session_state.transcriptions, diarization_enabled, include_speakers)
            total_tokens = count_tokens(text_for_llm, "HuggingFace" if token_backend.startswith("HuggingFace") else "tiktoken", hf_tokenizer_model)
            
            st.markdown(f"**Tokens**: `{total_tokens}` | **Segments**: `{len(st.session_state.transcriptions)}`")
            st.progress(min(1.0, total_tokens / max(1, context_limit)))
            
            # Vérifier si les transcriptions sont au format diarisation (tuple de 4)
            is_diarized = (
                st.session_state.transcriptions and 
                isinstance(st.session_state.transcriptions[0], (list, tuple)) and 
                len(st.session_state.transcriptions[0]) == 4
            )
            
            if is_diarized:
                with st.expander("✏️ Renommer locuteurs"):
                    with st.form("rename"):
                        for spk in sorted({item[2] for item in st.session_state.transcriptions}):
                            st.session_state.speaker_mapping[spk] = st.text_input(f"{spk} →", value=st.session_state.speaker_mapping.get(spk, spk), key=f"r_{spk}")
                        if st.form_submit_button("Appliquer"):
                            mapped = [(s, e, st.session_state.speaker_mapping.get(sp, sp), t) for s, e, sp, t in st.session_state.transcriptions]
                            st.session_state.transcriptions = mapped
                            st.rerun()
                
                for start, end, spk, txt in st.session_state.transcriptions:
                    st.markdown(f"- `{format_duration(start)}–{format_duration(end)}` **{spk}**: {txt}")
            else:
                # Sans diarisation - simple liste de textes
                for t in st.session_state.transcriptions:
                    st.markdown(f"- {t}")
    
    with tab3:
        col1, col2 = st.columns(2)
        
        # Vérifier le format des transcriptions
        is_diarized = (
            st.session_state.transcriptions and 
            isinstance(st.session_state.transcriptions[0], (list, tuple)) and 
            len(st.session_state.transcriptions[0]) == 4
        )
        
        with col1:
            if st.session_state.transcriptions:
                if is_diarized:
                    lines = [f"{format_duration(s)}-{format_duration(e)}: {sp}\n{t}" for s,e,sp,t in st.session_state.transcriptions]
                else:
                    lines = [str(t) for t in st.session_state.transcriptions]
                st.download_button("📄 Transcription", "\n\n".join(lines).encode(), "transcription.txt", use_container_width=True)
        with col2:
            if st.session_state.summary_text:
                st.download_button("📝 Résumé", st.session_state.summary_text.encode(), "resume.txt", use_container_width=True)
        
        if HAVE_DOCX and (st.session_state.summary_text or st.session_state.transcriptions):
            if st.button("📄 Générer DOCX", use_container_width=True):
                doc = Document()
                doc.add_heading("Compte rendu", 1)
                doc.add_paragraph(f"Date: {time.strftime('%d/%m/%Y %H:%M')}")
                if st.session_state.summary_text:
                    doc.add_heading("Résumé", 2)
                    doc.add_paragraph(st.session_state.summary_text)
                doc.add_heading("Transcription", 2)
                if is_diarized:
                    for s,e,sp,t in st.session_state.transcriptions:
                        doc.add_paragraph(f"{format_duration(s)}-{format_duration(e)} {sp}: {t}")
                else:
                    for t in st.session_state.transcriptions:
                        doc.add_paragraph(str(t))
                buf = io.BytesIO()
                doc.save(buf)
                st.download_button("⬇️ Télécharger", buf.getvalue(), "compte_rendu.docx", use_container_width=True)
else:
    st.info("👆 Chargez ou enregistrez un audio puis lancez le traitement.")

st.write("---")
st.caption("💡 L'enregistrement avancé génère un fichier WAV à télécharger puis importer.")